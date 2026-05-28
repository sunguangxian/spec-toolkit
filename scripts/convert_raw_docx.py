import argparse
from pathlib import Path
from typing import Iterable, List

from docx import Document

from atspec.core import ROOT, write_text


def escape_cell(text: str) -> str:
    return text.replace("\n", "<br>").replace("|", "\\|").strip()


def paragraph_text(paragraph) -> str:
    return "".join(run.text for run in paragraph.runs).strip()


def iter_block_items(doc: Document) -> Iterable[object]:
    body = doc.element.body
    paragraphs = {p._p: p for p in doc.paragraphs}
    tables = {t._tbl: t for t in doc.tables}
    for child in body.iterchildren():
        if child in paragraphs:
            yield paragraphs[child]
        elif child in tables:
            yield tables[child]


def table_to_markdown(table) -> List[str]:
    rows = []
    for row in table.rows:
        cells = []
        for cell in row.cells:
            parts = [p.text.strip() for p in cell.paragraphs if p.text.strip()]
            cells.append(escape_cell("\n".join(parts)))
        rows.append(cells)

    if not rows:
        return []

    width = max(len(row) for row in rows)
    normalized = [row + [""] * (width - len(row)) for row in rows]
    lines = []
    lines.append("| " + " | ".join(normalized[0]) + " |")
    lines.append("| " + " | ".join(["---"] * width) + " |")
    for row in normalized[1:]:
        lines.append("| " + " | ".join(row) + " |")
    return lines


def paragraph_to_markdown(paragraph) -> List[str]:
    text = paragraph_text(paragraph)
    if not text:
        return []

    style_name = (paragraph.style.name if paragraph.style is not None else "").lower()
    if "heading 1" in style_name or "标题 1" in style_name:
        return [f"# {text}"]
    if "heading 2" in style_name or "标题 2" in style_name:
        return [f"## {text}"]
    if "heading 3" in style_name or "标题 3" in style_name:
        return [f"### {text}"]
    if "heading 4" in style_name or "标题 4" in style_name:
        return [f"#### {text}"]
    return [text]


def convert_docx_to_markdown(input_path: Path, output_path: Path) -> None:
    doc = Document(str(input_path))
    lines: List[str] = []
    lines.append(f"<!-- Converted from: {input_path.as_posix()} -->")
    lines.append("")

    for block in iter_block_items(doc):
        if hasattr(block, "rows"):
            table_lines = table_to_markdown(block)
            if table_lines:
                lines.extend(table_lines)
                lines.append("")
        else:
            paragraph_lines = paragraph_to_markdown(block)
            if paragraph_lines:
                lines.extend(paragraph_lines)
                lines.append("")

    write_text(output_path, "\n".join(lines).rstrip() + "\n")


def find_docx_files(path: Path) -> List[Path]:
    if path.is_file():
        return [path]
    return sorted(p for p in path.rglob("*.docx") if not p.name.startswith("~$"))


def main() -> None:
    parser = argparse.ArgumentParser(description="Convert raw DOCX source files to Markdown for review and AT spec extraction.")
    parser.add_argument("--input", default="raw", help="DOCX file or directory. Default: raw")
    parser.add_argument("--output-dir", default="output/raw_converted", help="Directory for generated Markdown files.")
    parser.add_argument("--adjacent", action="store_true", help="Write Markdown next to each DOCX instead of output/raw_converted.")
    args = parser.parse_args()

    input_path = (ROOT / args.input).resolve() if not Path(args.input).is_absolute() else Path(args.input)
    output_dir = (ROOT / args.output_dir).resolve() if not Path(args.output_dir).is_absolute() else Path(args.output_dir)
    docx_files = find_docx_files(input_path)
    if not docx_files:
        raise SystemExit(f"No .docx files found: {input_path}")

    for docx_path in docx_files:
        if args.adjacent:
            output_path = docx_path.with_suffix(".md")
        else:
            relative = docx_path.relative_to(ROOT)
            output_path = output_dir / relative.with_suffix(".md")
        output_path.parent.mkdir(parents=True, exist_ok=True)
        convert_docx_to_markdown(docx_path, output_path)
        print(f"Converted: {docx_path.relative_to(ROOT)} -> {output_path.relative_to(ROOT)}")


if __name__ == "__main__":
    main()
